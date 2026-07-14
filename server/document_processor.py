"""
文档处理模块
支持PDF、Word、Markdown、TXT文件
"""
import os
import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict


class DocumentProcessor:
    """文档处理器"""

    def process_file(self, file_path: str) -> List[Dict]:
        """根据文件类型自动处理"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._process_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._process_word(file_path)
        elif ext == '.md':
            return self._process_markdown(file_path)
        elif ext == '.txt':
            return self._process_txt(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")

    def _process_pdf(self, file_path: str) -> List[Dict]:
        """处理PDF - 使用PyMuPDF"""
        doc = fitz.open(file_path)
        pages = []
        for i in range(len(doc)):
            page = doc[i]
            blocks = page.get_text("dict")["blocks"]
            lines = self._reconstruct_lines(blocks)
            text = "\n".join(lines)
            if text.strip():
                pages.append({
                    'text': text,
                    'metadata': {
                        'source': os.path.basename(file_path),
                        'page': i + 1,
                        'type': 'pdf'
                    }
                })
        doc.close()
        return pages

    def _reconstruct_lines(self, blocks: List[Dict]) -> List[str]:
        """重组文本行为逻辑行"""
        all_spans = []
        for block in blocks:
            if block["type"] == 0:
                for line in block["lines"]:
                    line_spans = []
                    for span in line["spans"]:
                        line_spans.append({
                            'text': span["text"],
                            'x0': span["bbox"][0],
                            'y0': span["bbox"][1]
                        })
                    if line_spans:
                        avg_y = sum(s['y0'] for s in line_spans) / len(line_spans)
                        all_spans.append({'spans': line_spans, 'y': avg_y})

        all_spans.sort(key=lambda x: x['y'])

        logical_lines = []
        current = []
        last_y = None
        for group in all_spans:
            if last_y is not None and abs(group['y'] - last_y) > 5:
                if current:
                    current.sort(key=lambda s: s['x0'])
                    logical_lines.append(" ".join(s['text'] for s in current))
                current = []
            current.extend(group['spans'])
            last_y = group['y']
        if current:
            current.sort(key=lambda s: s['x0'])
            logical_lines.append(" ".join(s['text'] for s in current))
        return logical_lines

    def _process_word(self, file_path: str) -> List[Dict]:
        """处理Word文档"""
        doc = Document(file_path)
        paragraphs = []
        current_text = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                current_text.append(text)
            elif current_text:
                paragraphs.append({
                    'text': "\n".join(current_text),
                    'metadata': {
                        'source': os.path.basename(file_path),
                        'type': 'word'
                    }
                })
                current_text = []

        if current_text:
            paragraphs.append({
                'text': "\n".join(current_text),
                'metadata': {
                    'source': os.path.basename(file_path),
                    'type': 'word'
                }
            })

        # 也处理表格
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                paragraphs.append({
                    'text': "\n".join(rows),
                    'metadata': {
                        'source': os.path.basename(file_path),
                        'type': 'word_table'
                    }
                })

        return paragraphs

    def _process_markdown(self, file_path: str) -> List[Dict]:
        """处理Markdown文件 - 按标题分割"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 按 ## 或 # 标题分割
        import re
        sections = re.split(r'\n(?=#{1,3}\s)', content)

        result = []
        for i, section in enumerate(sections):
            text = section.strip()
            if text:
                # 提取标题
                title_match = re.match(r'#{1,3}\s+(.+)', text)
                title = title_match.group(1) if title_match else f"段落{i + 1}"

                result.append({
                    'text': text,
                    'metadata': {
                        'source': os.path.basename(file_path),
                        'title': title,
                        'type': 'markdown'
                    }
                })

        if not result:
            # 如果没找到标题，整篇作为一个文档
            result.append({
                'text': content.strip(),
                'metadata': {
                    'source': os.path.basename(file_path),
                    'type': 'markdown'
                }
            })

        return result

    def _process_txt(self, file_path: str) -> List[Dict]:
        """处理TXT文件 - 按空行分段"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        paragraphs = []
        for para in content.split('\n\n'):
            text = para.strip()
            if text:
                paragraphs.append({
                    'text': text,
                    'metadata': {
                        'source': os.path.basename(file_path),
                        'type': 'txt'
                    }
                })

        if not paragraphs:
            paragraphs.append({
                'text': content.strip(),
                'metadata': {
                    'source': os.path.basename(file_path),
                    'type': 'txt'
                }
            })

        return paragraphs